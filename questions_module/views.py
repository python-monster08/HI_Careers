from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from django.utils.html import strip_tags
from .models import Question
import os
from datetime import datetime
import csv
from django.utils.timezone import now

def clean_text(text):
    """
    Replace newline characters with a single space and strip HTML tags.
    """
    return strip_tags(text.replace('\n', ' '))

def generate_questions_document(request):
    try:
        base_dir = r'F:/Kamlesh Lovewanshi/Office Work/generated_word_files'
        os.makedirs(base_dir, exist_ok=True)
        today = datetime.today().strftime('%Y-%m-%d')  # Format the date as YYYY-MM-DD
        file_name = f'all_questions_{today}.docx'  # Include the date in the filename
        file_path = os.path.join(base_dir, file_name)
        document = Document()

        for question in Question.objects.all():
            table = document.add_table(rows=0, cols=3)
            table.style = 'Table Grid'

            # Adding Question
            q_row = table.add_row().cells
            q_row[0].text = 'Question'
            q_row[1].text = clean_text(question.question)
            q_row[1].merge(q_row[2])

            # Adding Type
            type_row = table.add_row().cells
            type_row[0].text = 'Type'
            type_row[1].text = question.question_type.type_name  # Assuming QuestionType has a 'name' attribute
            type_row[1].merge(type_row[2])

            # Handling different question types
            if question.question_type.type_name.lower() == 'multiple_choice':
                options = [('A', question.option_a), ('B', question.option_b), 
                           ('C', question.option_c), ('D', question.option_d)]
                for label, option in options:
                    if option:  # Check if option is not None or blank
                        opt_row = table.add_row().cells
                        opt_row[0].text = 'Option'
                        opt_row[1].text = f"{label}: {clean_text(option)}"
                        opt_row[2].text = 'correct' if label == question.correct_answer_option else 'incorrect'

                # Solution
                solution_row = table.add_row().cells
                solution_row[0].text = 'Solution'
                solution_row[1].text = clean_text(question.correct_answer_description)
                solution_row[1].merge(solution_row[2])

            elif question.question_type.type_name.lower() in ['integer', 'true_false']:
                # Answer
                answer_row = table.add_row().cells
                answer_row[0].text = 'Answer'
                answer_row[1].text = clean_text(question.option_a)  # Using option_a for the answer
                answer_row[1].merge(answer_row[2])

                # Solution
                solution_row = table.add_row().cells
                solution_row[0].text = 'Solution'
                solution_row[1].text = clean_text(question.correct_answer_description)
                solution_row[1].merge(solution_row[2])

            # Marks
            marks_row = table.add_row().cells
            marks_row[0].text = 'Marks'
            marks_row[1].text = str(question.marks)
            marks_row[2].text = str(question.negative_marks) if question.negative_marks else "0"

            document.add_paragraph()  # Adds space between questions

        # Saving the document
        document.save(file_path)

        # Set success message
        message = f"Document created at: {file_path}"
        alert_type = 'success'

    except Exception as e:
        # Set error message
        message = f"An error occurred: {e}"
        alert_type = 'danger'

    context = {'message': message, 'alert_type': alert_type}
    return render(request, 'content_manager/files_genrater_or_uploader.html', context)




def generate_questions_csv_view(request):
    # Define the directory and filename
    base_dir = 'F:/Kamlesh Lovewanshi/Office Work/generated_word_files'
    os.makedirs(base_dir, exist_ok=True)  # Ensure the directory exists
    file_name = f'questions_{now().strftime("%Y-%m-%d")}.csv'
    file_path = os.path.join(base_dir, file_name)

    try:
        with open(file_path, mode='w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            # Write the header
            writer.writerow(['SrNo.', 'Question', 'Option a', 'Option b', 'Option c', 'Option d', 'Correct Option(a/b/c/d)', 'Keywords'])

            for index, question in enumerate(Question.objects.all(), start=1):
                # Retrieve keywords
                keywords = ', '.join(keyword.name for keyword in question.keywords.all())

                writer.writerow([
                    index,  # SrNo.
                    question.question,  # Question
                    question.option_a,  # Option a
                    question.option_b or '',  # Option b
                    question.option_c or '',  # Option c
                    question.option_d or '',  # Option d
                    question.correct_answer_option.upper(),  # Correct Option
                    keywords or ''  # Keywords
                ])

        # Set success message
        message = f"CSV file has been successfully generated at: {file_path}"
        alert_type = 'success'

    except PermissionError as e:
        # Set error message
        message = f"Permission error: {e}"
        alert_type = 'danger'

    context = {'message': message, 'alert_type': alert_type}
    return render(request, 'content_manager/files_genrater_or_uploader.html', context)



    
def index(request):
    return render(request, 'index.html')